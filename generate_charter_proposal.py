#!/usr/bin/env python3
"""
FS Vector Charter Proposal Generator

Generates bank charter advisory proposals as .docx files by cloning the
Onyx charter template and replacing variable content sections.
"""

import copy
import os
import re
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Reuse utilities from the regulatory proposal generator
from generate_proposal import (
    fetch_company_logo,
    insert_logo_into_paragraph,
    build_toc_sdt,
    set_updatefields_on_open,
    _strip_includepicture_fields,
    clear_paragraph_runs,
    replace_client_name_in_paragraph,
    remove_element,
)

# ── Template path ─────────────────────────────────────────────────────
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_TEMPLATE_NAME = "FS Vector Charter Proposal Template.docx"

CHARTER_TEMPLATE_PATH = os.path.join(_SCRIPT_DIR, _TEMPLATE_NAME)
if not os.path.exists(CHARTER_TEMPLATE_PATH):
    CHARTER_TEMPLATE_PATH = os.path.expanduser(f"~/Desktop/{_TEMPLATE_NAME}")


# =====================================================================
# Paragraph creation helpers (charter-specific formatting)
# =====================================================================

def _create_bold_heading_para(text, template_para):
    """Create a bold heading paragraph (e.g., 'Key Deliverables, Activities, and Outcomes')."""
    new_elem = copy.deepcopy(template_para._element)
    # Clear all runs
    for r in new_elem.findall(qn('w:r')):
        new_elem.remove(r)
    # Remove pageBreakBefore so cloned paras don't create blank pages
    pPr = new_elem.find(qn('w:pPr'))
    if pPr is not None:
        pbb = pPr.find(qn('w:pageBreakBefore'))
        if pbb is not None:
            pPr.remove(pbb)
    # Add single bold run
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Roboto')
    rFonts.set(qn('w:hAnsi'), 'Roboto')
    rPr.append(rFonts)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    new_elem.append(run)
    return new_elem


def _create_deliverable_para(name, description, template_para):
    """Create a deliverable paragraph with bold lead-in name and regular description.

    Format: **Name**: Description text...
    With left indent of 720 twips (0.5 inches).
    """
    new_elem = copy.deepcopy(template_para._element)
    # Clear all runs
    for r in new_elem.findall(qn('w:r')):
        new_elem.remove(r)

    # Set paragraph indent
    pPr = new_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        new_elem.insert(0, pPr)
    # Remove pageBreakBefore
    pbb = pPr.find(qn('w:pageBreakBefore'))
    if pbb is not None:
        pPr.remove(pbb)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), '720')

    # Bold run for the name
    run1 = OxmlElement('w:r')
    rPr1 = OxmlElement('w:rPr')
    b1 = OxmlElement('w:b')
    rPr1.append(b1)
    rFonts1 = OxmlElement('w:rFonts')
    rFonts1.set(qn('w:ascii'), 'Roboto')
    rFonts1.set(qn('w:hAnsi'), 'Roboto')
    rPr1.append(rFonts1)
    run1.append(rPr1)
    t1 = OxmlElement('w:t')
    t1.text = name
    t1.set(qn('xml:space'), 'preserve')
    run1.append(t1)
    new_elem.append(run1)

    # Regular run for ": description"
    run2 = OxmlElement('w:r')
    rPr2 = OxmlElement('w:rPr')
    rFonts2 = OxmlElement('w:rFonts')
    rFonts2.set(qn('w:ascii'), 'Roboto')
    rFonts2.set(qn('w:hAnsi'), 'Roboto')
    rPr2.append(rFonts2)
    run2.append(rPr2)
    t2 = OxmlElement('w:t')
    t2.text = ": " + description
    t2.set(qn('xml:space'), 'preserve')
    run2.append(t2)
    new_elem.append(run2)

    return new_elem


def _create_sub_bullet_para(text, num_id):
    """Create a sub-bullet paragraph matching the charter template's numbered list style.

    indent: left=1440, hanging=360
    numPr: ilvl=0, numId from template
    """
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    # Numbering
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numPr.append(ilvl)
    nid = OxmlElement('w:numId')
    nid.set(qn('w:val'), str(num_id))
    numPr.append(nid)
    pPr.append(numPr)

    # Indent
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '1440')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)

    p.append(pPr)

    # Text run
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Roboto')
    rFonts.set(qn('w:hAnsi'), 'Roboto')
    rPr.append(rFonts)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    p.append(run)

    return p


def _create_body_para(text, template_para):
    """Create a normal body paragraph."""
    new_elem = copy.deepcopy(template_para._element)
    for r in new_elem.findall(qn('w:r')):
        new_elem.remove(r)
    # Remove any indent, numbering, and page break before
    pPr = new_elem.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
        # Remove pageBreakBefore so cloned paras don't create blank pages
        pbb = pPr.find(qn('w:pageBreakBefore'))
        if pbb is not None:
            pPr.remove(pbb)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Roboto')
    rFonts.set(qn('w:hAnsi'), 'Roboto')
    rPr.append(rFonts)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    new_elem.append(run)
    return new_elem


def _create_fee_heading_para(template_para):
    """Create the bold 'Fee and Timeline' heading."""
    return _create_bold_heading_para("Fee and Timeline", template_para)


def _create_fee_text_para(text, template_para):
    """Create a fee description paragraph. The full text is regular weight
    (fee amounts will naturally appear in the AI-generated text)."""
    new_elem = copy.deepcopy(template_para._element)
    for r in new_elem.findall(qn('w:r')):
        new_elem.remove(r)
    pPr = new_elem.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)
        pbb = pPr.find(qn('w:pageBreakBefore'))
        if pbb is not None:
            pPr.remove(pbb)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Roboto')
    rFonts.set(qn('w:hAnsi'), 'Roboto')
    rPr.append(rFonts)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    new_elem.append(run)
    return new_elem


def _create_heading2_para(text, template_heading):
    """Create a Heading 2 paragraph by cloning an existing one."""
    new_elem = copy.deepcopy(template_heading._element)
    for r in new_elem.findall(qn('w:r')):
        new_elem.remove(r)
    run = OxmlElement('w:r')
    rPr_source = template_heading.runs[0]._element.find(qn('w:rPr')) if template_heading.runs else None
    if rPr_source is not None:
        run.append(copy.deepcopy(rPr_source))
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    new_elem.append(run)
    return new_elem


# =====================================================================
# Main generator
# =====================================================================

def generate_charter_proposal(
    client_name,
    proposal_type,
    date_str,
    engagement_background_paragraphs,
    engagement_scope_phases,
    client_website=None,
    output_path=None,
):
    """
    Generate a bank charter advisory proposal document.

    Args:
        client_name: Name of the client company
        proposal_type: Title line (e.g., "Bank Charter Advisory Services")
        date_str: Date string (e.g., "February 2026")
        engagement_background_paragraphs: List of strings for Engagement Background
        engagement_scope_phases: List of phase dicts, each with:
            - "heading": phase title (Heading 2)
            - "intro": intro paragraph text
            - "deliverables": list of dicts with "name", "description", "sub_items"
            - "fee_text": fee and timeline paragraph
            - "bridge_text": optional transition paragraph to next phase
        client_website: URL for logo fetching (optional)
        output_path: Where to save. If None, auto-generates path.

    Returns:
        Path to the generated document.
    """
    if not os.path.exists(CHARTER_TEMPLATE_PATH):
        raise FileNotFoundError(f"Charter template not found at: {CHARTER_TEMPLATE_PATH}")

    doc = Document(CHARTER_TEMPLATE_PATH)
    paragraphs = doc.paragraphs

    # ===== 1. UPDATE COVER PAGE =====
    # P3: "Proposal to Provide" — keep as-is
    # P4: proposal type title
    if len(paragraphs) > 4:
        clear_paragraph_runs(paragraphs[4])
        run = paragraphs[4].add_run(proposal_type)
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x07, 0x37, 0x63)

    # P9: Client logo
    if client_website:
        logo_path = fetch_company_logo(client_website)
        if logo_path:
            # Clear existing content in P9
            clear_paragraph_runs(paragraphs[9])
            for drawing in paragraphs[9]._element.findall('.//' + qn('wp:inline')):
                drawing.getparent().getparent().getparent().remove(
                    drawing.getparent().getparent()
                )
            insert_logo_into_paragraph(paragraphs[9], logo_path, width_inches=2.0)
    else:
        # Replace logo with client name text
        clear_paragraph_runs(paragraphs[9])
        for drawing in paragraphs[9]._element.findall('.//' + qn('wp:inline')):
            drawing.getparent().getparent().getparent().remove(
                drawing.getparent().getparent()
            )
        run = paragraphs[9].add_run(client_name)
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x07, 0x37, 0x63)

    # ===== 2. UPDATE HEADERS =====
    for section in doc.sections:
        header = section.header
        if header and header.paragraphs:
            for hp in header.paragraphs:
                if 'Proposal to Provide' in hp.text or 'Bank Charter' in hp.text:
                    for run in hp.runs:
                        if 'Proposal to Provide' in run.text or 'Bank Charter' in run.text:
                            run.text = f"Proposal to Provide {proposal_type}"
                if any(month in hp.text for month in [
                    'January', 'February', 'March', 'April', 'May', 'June',
                    'July', 'August', 'September', 'October', 'November', 'December'
                ]):
                    for run in hp.runs:
                        run.text = date_str

    # ===== 3. REPLACE TOC =====
    # The template has: page-break paragraph, "Contents", SDT, empty paragraphs,
    # page-break paragraph. We remove all of that and insert fresh page breaks
    # around a new TOC SDT so the TOC gets its own page.
    body = doc.element.body

    # First, remove ALL existing SDT elements (the template's old TOC)
    for sdt in body.findall(qn('w:sdt')):
        sdt.getparent().remove(sdt)

    # Find "Contents" paragraph and the surrounding area (including page breaks)
    toc_title_idx = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == 'Contents':
            toc_title_idx = i
            break

    if toc_title_idx is not None:
        toc_title_para = paragraphs[toc_title_idx]

        # Scan backward from "Contents" to find page-break paragraphs
        toc_start_idx = toc_title_idx
        for i in range(toc_title_idx - 1, max(toc_title_idx - 5, -1), -1):
            p = paragraphs[i]
            is_empty_or_pagebreak = (p.text.strip() == '')
            if is_empty_or_pagebreak:
                toc_start_idx = i
            else:
                break

        # Scan forward from "Contents" to find trailing empty/page-break paragraphs
        toc_end_idx = toc_title_idx
        for i in range(toc_title_idx + 1, min(toc_title_idx + 10, len(paragraphs))):
            if paragraphs[i].text.strip() == '':
                toc_end_idx = i
            else:
                break

        # Build a fresh TOC field
        toc_sdt = build_toc_sdt(toc_title_format_para=toc_title_para)

        # Create page break paragraphs to isolate the TOC on its own page
        pb_before = OxmlElement('w:p')
        pb_before_r = OxmlElement('w:r')
        pb_before_br = OxmlElement('w:br')
        pb_before_br.set(qn('w:type'), 'page')
        pb_before_r.append(pb_before_br)
        pb_before.append(pb_before_r)

        pb_after = OxmlElement('w:p')
        pb_after_r = OxmlElement('w:r')
        pb_after_br = OxmlElement('w:br')
        pb_after_br.set(qn('w:type'), 'page')
        pb_after_r.append(pb_after_br)
        pb_after.append(pb_after_r)

        # Insert: page break, TOC SDT, page break — before the old content
        anchor = paragraphs[toc_start_idx]._element
        anchor.addprevious(pb_before)
        anchor.addprevious(toc_sdt)
        anchor.addprevious(pb_after)

        # Remove the old "Contents" heading + surrounding empties + page breaks
        for i in range(toc_start_idx, toc_end_idx + 1):
            try:
                paragraphs[i]._element.getparent().remove(paragraphs[i]._element)
            except (ValueError, AttributeError):
                pass

    set_updatefields_on_open(doc)

    # ===== 4. REPLACE ENGAGEMENT BACKGROUND =====
    # Find "Engagement Background" heading
    bg_start = None
    bg_end = None
    for i, p in enumerate(paragraphs):
        if p.style and 'Heading 1' in p.style.name and p.text.strip() == 'Engagement Background':
            bg_start = i
        elif bg_start is not None and p.style and 'Heading 1' in p.style.name and p.text.strip():
            bg_end = i
            break

    if bg_start is not None and bg_end is not None:
        # Get a template body paragraph for formatting reference
        template_body = None
        for i in range(bg_start + 1, bg_end):
            if paragraphs[i].text.strip():
                template_body = paragraphs[i]
                break

        if template_body is None:
            template_body = paragraphs[bg_start + 1]

        # Remove old background paragraphs (keep the heading)
        elements_to_remove = []
        for i in range(bg_start + 1, bg_end):
            elements_to_remove.append(paragraphs[i]._element)

        # Insert new background paragraphs
        insert_after = paragraphs[bg_start]._element
        for para_text in engagement_background_paragraphs:
            new_elem = _create_body_para(para_text, template_body)
            insert_after.addnext(new_elem)
            insert_after = new_elem

        # Remove old elements
        for elem in elements_to_remove:
            elem.getparent().remove(elem)

    # ===== 5. REPLACE ENGAGEMENT SCOPE PHASES =====
    # Find "Engagement Scope and Fees" heading
    scope_start = None
    scope_end = None
    # Refresh paragraphs after modifications
    paragraphs = doc.paragraphs

    for i, p in enumerate(paragraphs):
        if p.style and 'Heading 1' in p.style.name and 'Engagement Scope' in p.text:
            scope_start = i
        elif scope_start is not None and p.style and 'Heading 1' in p.style.name and p.text.strip() and 'Engagement Scope' not in p.text:
            scope_end = i
            break

    if scope_start is not None and scope_end is not None:
        # Find a Heading 2 template for phase headings
        heading2_template = None
        body_template = None
        for i in range(scope_start, scope_end):
            if paragraphs[i].style and 'Heading 2' in paragraphs[i].style.name and heading2_template is None:
                heading2_template = paragraphs[i]
            if paragraphs[i].text.strip() and paragraphs[i].style and paragraphs[i].style.name == 'normal' and body_template is None:
                body_template = paragraphs[i]

        # Find the numId used for sub-bullets in the template.
        # IMPORTANT: skip Heading 2 paragraphs (they use numId=10 for
        # outline numbering) — we only want the bullet list numId (=8).
        num_id = 8  # default from template analysis
        for i in range(scope_start, scope_end):
            p_style = paragraphs[i].style
            if p_style and 'Heading' in p_style.name:
                continue  # skip headings — they have their own numbering
            numPr = paragraphs[i]._element.find('.//' + qn('w:numPr'))
            if numPr is not None:
                nid_elem = numPr.find(qn('w:numId'))
                if nid_elem is not None:
                    num_id = int(nid_elem.get(qn('w:val')))
                    break

        if body_template is None:
            body_template = paragraphs[scope_start + 1]

        # Remove old scope content (keep the "Engagement Scope and Fees" heading)
        # Also keep the intro paragraph right after it
        scope_intro_end = scope_start + 1
        elements_to_remove = []
        for i in range(scope_start + 1, scope_end):
            elements_to_remove.append(paragraphs[i]._element)

        # Build new scope content
        insert_after = paragraphs[scope_start]._element

        # Scope intro paragraph
        scope_intro_text = (
            f"FS Vector will structure its support for {client_name} across the "
            f"{_num_to_words(len(engagement_scope_phases))} phases described below. "
            f"The phases described here constitute FS Vector's recommended scope "
            f"for the engagement."
        )
        intro_elem = _create_body_para(scope_intro_text, body_template)
        insert_after.addnext(intro_elem)
        insert_after = intro_elem

        # Insert each phase
        for phase in engagement_scope_phases:
            # Phase heading (Heading 2)
            if heading2_template is not None:
                h2_elem = _create_heading2_para(phase["heading"], heading2_template)
            else:
                h2_elem = _create_bold_heading_para(phase["heading"], body_template)
            insert_after.addnext(h2_elem)
            insert_after = h2_elem

            # Phase intro paragraph
            if phase.get("intro"):
                intro_elem = _create_body_para(phase["intro"], body_template)
                insert_after.addnext(intro_elem)
                insert_after = intro_elem

            # "Key Deliverables, Activities, and Outcomes" heading
            kd_elem = _create_bold_heading_para(
                "Key Deliverables, Activities, and Outcomes", body_template
            )
            insert_after.addnext(kd_elem)
            insert_after = kd_elem

            # Deliverables
            for deliv in phase.get("deliverables", []):
                # Bold lead-in + description
                d_elem = _create_deliverable_para(
                    deliv["name"], deliv.get("description", ""), body_template
                )
                insert_after.addnext(d_elem)
                insert_after = d_elem

                # Sub-items as bullets
                for sub in deliv.get("sub_items", []):
                    sb_elem = _create_sub_bullet_para(sub, num_id)
                    insert_after.addnext(sb_elem)
                    insert_after = sb_elem

            # Fee and Timeline heading
            fee_h_elem = _create_fee_heading_para(body_template)
            insert_after.addnext(fee_h_elem)
            insert_after = fee_h_elem

            # Fee text
            if phase.get("fee_text"):
                fee_elem = _create_fee_text_para(phase["fee_text"], body_template)
                insert_after.addnext(fee_elem)
                insert_after = fee_elem

            # Bridge text (transition to next phase)
            if phase.get("bridge_text"):
                bridge_elem = _create_body_para(phase["bridge_text"], body_template)
                insert_after.addnext(bridge_elem)
                insert_after = bridge_elem

        # Remove old elements
        for elem in elements_to_remove:
            try:
                elem.getparent().remove(elem)
            except (ValueError, AttributeError):
                pass

    # ===== 6. REPLACE CLIENT NAME THROUGHOUT =====
    paragraphs = doc.paragraphs
    for p in paragraphs:
        replace_client_name_in_paragraph(p, "Onyx", client_name)

    # Also replace in headers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                replace_client_name_in_paragraph(p, "Onyx", client_name)

    # ===== 7. FIX BLANK PAGE BREAK PARAGRAPHS =====
    # The template uses empty paragraphs with page breaks as section separators.
    # These often carry formatting (Heading1 style with large spacing, bold fonts,
    # spacing-before, keepNext, etc.) that causes Word to render visible blank
    # space before the break, creating unwanted blank pages. Replace each one
    # with a clean, minimal page-break-only paragraph.
    paragraphs = doc.paragraphs
    for p in paragraphs:
        if p.text.strip() != '':
            continue
        has_page_break = False
        for run_elem in p._element.findall(qn('w:r')):
            for br in run_elem.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    has_page_break = True
        if not has_page_break:
            continue
        # Replace the entire paragraph content with a clean page break
        elem = p._element
        # Remove all existing children
        for child in list(elem):
            elem.remove(child)
        # Add a single clean run with just the page break
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        elem.append(r)

    # ===== 8. STRIP INCLUDEPICTURE FIELDS =====
    _strip_includepicture_fields(doc)

    # ===== 8. SAVE =====
    if output_path is None:
        safe_name = client_name.replace(" ", "_")
        output_path = os.path.join(
            os.path.expanduser("~/Desktop"),
            f"FS_Vector_Charter_Proposal_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx",
        )

    doc.save(output_path)
    return output_path


def _num_to_words(n):
    """Convert small integers to words."""
    words = {1: "one", 2: "two", 3: "three", 4: "four", 5: "five",
             6: "six", 7: "seven", 8: "eight", 9: "nine", 10: "ten"}
    return words.get(n, str(n))
