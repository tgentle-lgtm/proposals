#!/usr/bin/env python3
"""
FS Vector Licensing Proposal Generator

Generates licensing services proposals as .docx files by cloning the
ether.fi licensing template and replacing variable content sections.

Licensing proposals are ~90% boilerplate — only the Engagement Background
section is AI-generated. All scope, fee tables, and timeline content remain
static from the template.
"""

import copy
import os
import re
import tempfile
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Reuse utilities from the regulatory proposal generator
from generate_proposal import (
    fetch_company_logo,
    insert_logo_into_paragraph,
    set_updatefields_on_open,
    _strip_includepicture_fields,
    clear_paragraph_runs,
    replace_client_name_in_paragraph,
    remove_element,
)

# ── Template path ─────────────────────────────────────────────────────
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_TEMPLATE_NAME = "FS Vector Licensing Proposal Template.docx"

LICENSING_TEMPLATE_PATH = os.path.join(_SCRIPT_DIR, _TEMPLATE_NAME)
if not os.path.exists(LICENSING_TEMPLATE_PATH):
    LICENSING_TEMPLATE_PATH = os.path.expanduser(f"~/Desktop/{_TEMPLATE_NAME}")

# Old client name in the template (to be replaced throughout)
_OLD_CLIENT_NAME = "ether.fi"


# =====================================================================
# Paragraph creation helper (licensing-specific formatting)
# =====================================================================

def _create_body_para(text, template_para):
    """Create a normal body paragraph with Roboto Light 11pt font."""
    new_elem = copy.deepcopy(template_para._element)
    for r in new_elem.findall(qn('w:r')):
        new_elem.remove(r)
    # Remove any indent, numbering, and page break before
    pPr = new_elem.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
        pbb = pPr.find(qn('w:pageBreakBefore'))
        if pbb is not None:
            pPr.remove(pbb)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Roboto Light')
    rFonts.set(qn('w:hAnsi'), 'Roboto Light')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')  # 11pt
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '22')
    rPr.append(szCs)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    new_elem.append(run)
    return new_elem


# =====================================================================
# Main generator
# =====================================================================

def generate_licensing_proposal(
    client_name,
    date_str,
    engagement_background_paragraphs,
    client_website=None,
    acquisition_fee=None,
    maintenance_fee=None,
    output_path=None,
):
    """
    Generate a licensing services proposal document.

    Args:
        client_name: Name of the client company
        date_str: Date string (e.g., "February 2026")
        engagement_background_paragraphs: List of strings for Engagement Background
        client_website: URL for logo fetching (optional)
        acquisition_fee: Monthly fee for license acquisition (e.g., "$30,000/month").
            If provided, replaces all phase fee cells in Table 1.
        maintenance_fee: Annual fee for license maintenance (e.g., "$180,000/year").
            If provided, replaces the fee cell in Table 2.
        output_path: Where to save. If None, auto-generates path.

    Returns:
        Path to the generated document.
    """
    if not os.path.exists(LICENSING_TEMPLATE_PATH):
        raise FileNotFoundError(
            f"Licensing template not found at: {LICENSING_TEMPLATE_PATH}"
        )

    doc = Document(LICENSING_TEMPLATE_PATH)
    paragraphs = doc.paragraphs

    # ===== 1. UPDATE COVER PAGE =====
    # P2: title ("Proposal to Provide Licensing Services") — keep as-is
    # P8: "Prepared for ether.fi" — replace with client name
    if len(paragraphs) > 8:
        clear_paragraph_runs(paragraphs[8])
        # Remove any existing images/drawings
        for drawing in paragraphs[8]._element.findall('.//' + qn('wp:inline')):
            drawing.getparent().getparent().getparent().remove(
                drawing.getparent().getparent()
            )
        run = paragraphs[8].add_run(f"Prepared for {client_name}")
        run.font.name = 'Roboto'
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x07, 0x37, 0x63)

    # P13: date
    if len(paragraphs) > 13:
        clear_paragraph_runs(paragraphs[13])
        run = paragraphs[13].add_run(date_str)
        run.font.name = 'Roboto'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x07, 0x37, 0x63)

    # ===== 2. CLIENT LOGO ON COVER PAGE =====
    # Try to insert logo into the cover page area (P9 or nearby)
    if client_website:
        logo_path = fetch_company_logo(client_website)
        if logo_path:
            # Find a suitable paragraph for the logo (after client name, before date)
            # Use P9 if available and empty, or insert into P8 area
            for logo_idx in [9, 10]:
                if logo_idx < len(paragraphs) and not paragraphs[logo_idx].text.strip():
                    insert_logo_into_paragraph(
                        paragraphs[logo_idx], logo_path, width_inches=2.0
                    )
                    break

    # ===== 3. UPDATE HEADERS =====
    for section in doc.sections:
        header = section.header
        if header and header.paragraphs:
            for hp in header.paragraphs:
                # Replace old client name in headers
                if _OLD_CLIENT_NAME in hp.text:
                    for run in hp.runs:
                        if _OLD_CLIENT_NAME in run.text:
                            run.text = run.text.replace(
                                _OLD_CLIENT_NAME, client_name
                            )
                # Replace date in headers
                if any(month in hp.text for month in [
                    'January', 'February', 'March', 'April', 'May', 'June',
                    'July', 'August', 'September', 'October', 'November',
                    'December'
                ]):
                    # Only replace if this looks like a date line (short text)
                    text = hp.text.strip()
                    if len(text) < 50:
                        for run in hp.runs:
                            run.text = date_str

    # ===== 4. REPLACE ENGAGEMENT BACKGROUND =====
    # Find "Engagement Background" heading
    bg_start = None
    bg_end = None
    for i, p in enumerate(paragraphs):
        if (p.style and 'Heading 1' in p.style.name
                and p.text.strip() == 'Engagement Background'):
            bg_start = i
        elif (bg_start is not None and p.style
              and 'Heading 1' in p.style.name and p.text.strip()):
            bg_end = i
            break

    if bg_start is not None and bg_end is not None:
        # Get a template body paragraph for formatting reference
        template_body = None
        for i in range(bg_start + 1, bg_end):
            if paragraphs[i].text.strip():
                template_body = paragraphs[i]
                break

        if template_body is None:
            template_body = paragraphs[bg_start + 1]

        # Remove old background paragraphs (keep the heading)
        elements_to_remove = []
        for i in range(bg_start + 1, bg_end):
            elements_to_remove.append(paragraphs[i]._element)

        # Insert new background paragraphs
        insert_after = paragraphs[bg_start]._element
        for para_text in engagement_background_paragraphs:
            new_elem = _create_body_para(para_text, template_body)
            insert_after.addnext(new_elem)
            insert_after = new_elem

        # Remove old elements
        for elem in elements_to_remove:
            elem.getparent().remove(elem)

    # ===== 5. REPLACE CLIENT NAME THROUGHOUT =====
    # The template uses both "ether.fi" and "Ether.fi" — replace both
    _old_names = [_OLD_CLIENT_NAME, _OLD_CLIENT_NAME.capitalize(), "Ether.fi"]
    # Deduplicate while preserving order
    old_names = list(dict.fromkeys(_old_names))

    # Refresh paragraphs after modifications
    paragraphs = doc.paragraphs
    for old in old_names:
        for p in paragraphs:
            replace_client_name_in_paragraph(p, old, client_name)

    # Also replace in headers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                for old in old_names:
                    replace_client_name_in_paragraph(p, old, client_name)

    # Replace in ALL table cells (ether.fi appears in fee/timeline tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old in old_names:
                        replace_client_name_in_paragraph(p, old, client_name)

    # ===== 6. UPDATE FEE TABLES =====
    # Table 1 = License Acquisition (Row 2 = "Monthly Fee", Cells 1-3 = phase fees)
    # Table 2 = License Maintenance (Row 2 = annual fee, merged cell)
    if acquisition_fee and len(doc.tables) > 1:
        acq_table = doc.tables[1]
        if len(acq_table.rows) > 2:
            fee_row = acq_table.rows[2]
            # Replace cells 1, 2, 3 (skip cell 0 which is the "Monthly Fee" label)
            for ci in range(1, min(4, len(fee_row.cells))):
                cell = fee_row.cells[ci]
                for p in cell.paragraphs:
                    clear_paragraph_runs(p)
                    run = p.add_run(acquisition_fee)
                    run.font.name = 'Roboto'
                    run.font.size = Pt(9)

    if maintenance_fee and len(doc.tables) > 2:
        maint_table = doc.tables[2]
        if len(maint_table.rows) > 2:
            fee_row = maint_table.rows[2]
            # Cell 1 is the merged fee cell (gridSpan=6); updating it covers all
            # Deduplicate cell refs since python-docx exposes merged cells multiple times
            seen_tcs = set()
            for ci in range(1, len(fee_row.cells)):
                cell = fee_row.cells[ci]
                tc_id = id(cell._tc)
                if tc_id in seen_tcs:
                    continue  # Skip duplicate refs to same merged cell
                seen_tcs.add(tc_id)
                # Write fee to first paragraph only, remove extras
                paras = cell.paragraphs
                if paras:
                    clear_paragraph_runs(paras[0])
                    run = paras[0].add_run(maintenance_fee)
                    run.font.name = 'Roboto'
                    run.font.size = Pt(9)
                    # Remove extra paragraphs from the merged cell
                    for extra_p in paras[1:]:
                        extra_p._element.getparent().remove(extra_p._element)

    # ===== 7. FIX BLANK PAGE BREAK PARAGRAPHS =====
    paragraphs = doc.paragraphs
    for p in paragraphs:
        if p.text.strip() != '':
            continue
        has_page_break = False
        for run_elem in p._element.findall(qn('w:r')):
            for br in run_elem.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    has_page_break = True
        if not has_page_break:
            continue
        # Replace the entire paragraph content with a clean page break
        elem = p._element
        for child in list(elem):
            elem.remove(child)
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        elem.append(r)

    # ===== 8. SET UPDATE FIELDS ON OPEN =====
    set_updatefields_on_open(doc)

    # ===== 9. STRIP INCLUDEPICTURE FIELDS =====
    _strip_includepicture_fields(doc)

    # ===== 10. SAVE =====
    if output_path is None:
        safe_name = client_name.replace(" ", "_")
        output_path = os.path.join(
            os.path.expanduser("~/Desktop"),
            f"FS_Vector_Licensing_Proposal_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx",
        )

    doc.save(output_path)
    return output_path
