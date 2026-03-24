#!/usr/bin/env python3
"""
FS Vector Proposal Generator

Generates client proposals as .docx files by cloning the Parent template
and replacing variable content sections while preserving all formatting.

Usage:
    python3 generate_proposal.py --interactive

The script will prompt for:
    - Client name and website URL (for logo fetching)
    - Proposal type (title)
    - Date
    - Engagement Background content (paragraphs)
    - Engagement Scope subsections
    - Budget and Timeline details
    - Whether to include Managed Services as a primary service
"""

import copy
import os
import sys
import re
import tempfile
from datetime import datetime
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Look for template in same directory as this script first, then parent, then Desktop
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_TEMPLATE_NAME = "FS Vector Parent Proposal Template.docx"

TEMPLATE_PATH = os.path.join(_SCRIPT_DIR, _TEMPLATE_NAME)
if not os.path.exists(TEMPLATE_PATH):
    TEMPLATE_PATH = os.path.join(_SCRIPT_DIR, "..", _TEMPLATE_NAME)
if not os.path.exists(TEMPLATE_PATH):
    TEMPLATE_PATH = os.path.expanduser(f"~/Desktop/{_TEMPLATE_NAME}")


# =====================================================================
# Logo Fetching
# =====================================================================

def fetch_company_logo(url):
    """Fetch a company logo from their website URL.

    Tries multiple strategies: og:image, img tags with 'logo', apple-touch-icon,
    high-res favicon. Downloads to a temp file and converts to PNG if needed.

    Returns path to a temporary PNG/JPG file, or None on failure.
    """
    headers = {
        'User-Agent': ('Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) '
                       'AppleWebKit/537.36 (KHTML, like Gecko) '
                       'Chrome/120.0.0.0 Safari/537.36')
    }

    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url

    try:
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
    except requests.RequestException as e:
        print(f"  Warning: Could not fetch website {url}: {e}")
        return None

    soup = BeautifulSoup(response.text, 'html.parser')
    logo_url = None

    # Strategy 1: og:image meta tag
    og_image = soup.find("meta", {"property": "og:image"})
    if og_image and og_image.get("content"):
        logo_url = og_image["content"]

    # Strategy 2: <img> tag with "logo" in alt, class, id, or src
    if not logo_url:
        for img in soup.find_all('img'):
            alt = (img.get('alt') or '').lower()
            cls = ' '.join(img.get('class') or []).lower()
            img_id = (img.get('id') or '').lower()
            src = (img.get('src') or '').lower()
            if any('logo' in attr for attr in [alt, cls, img_id, src]):
                logo_url = img.get('src')
                break

    # Strategy 3: apple-touch-icon (usually high quality)
    if not logo_url:
        apple_icon = soup.find('link', rel=lambda r: r and 'apple-touch-icon' in r)
        if apple_icon and apple_icon.get('href'):
            logo_url = apple_icon['href']

    # Strategy 4: High-res favicon link
    if not logo_url:
        for link in soup.find_all('link'):
            rel = ' '.join(link.get('rel', [])).lower()
            if 'icon' in rel and 'shortcut' not in rel:
                # Prefer larger icons
                sizes = link.get('sizes', '')
                logo_url = link.get('href')
                if sizes and any(int(s.split('x')[0]) >= 128
                                 for s in sizes.split() if 'x' in s):
                    break

    # Strategy 5: Fallback to /favicon.ico
    if not logo_url:
        logo_url = urljoin(url, '/favicon.ico')

    # Resolve relative URLs
    if logo_url and not logo_url.startswith(('http://', 'https://')):
        logo_url = urljoin(url, logo_url)

    # Download the image
    if logo_url:
        return _download_and_convert_logo(logo_url, headers)

    return None


def _download_and_convert_logo(image_url, headers):
    """Download an image and convert to PNG if needed. Returns file path or None."""
    try:
        resp = requests.get(image_url, headers=headers, timeout=15)
        if resp.status_code != 200 or len(resp.content) < 100:
            return None

        content_type = resp.headers.get('Content-Type', '').lower()

        # Determine if conversion is needed
        needs_conversion = any(fmt in content_type for fmt in ['svg', 'ico', 'webp'])
        if not needs_conversion and image_url.lower().endswith(('.svg', '.ico', '.webp')):
            needs_conversion = True

        if needs_conversion:
            return _convert_to_png(resp.content)

        # Already a usable format (PNG, JPG)
        ext = '.png'
        if 'jpeg' in content_type or 'jpg' in content_type:
            ext = '.jpg'

        tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
        tmp.write(resp.content)
        tmp.close()
        return tmp.name

    except requests.RequestException as e:
        print(f"  Warning: Could not download logo from {image_url}: {e}")
        return None


def _convert_to_png(image_bytes):
    """Convert image bytes (SVG, ICO, WebP, etc.) to PNG using Pillow."""
    try:
        from PIL import Image
        import io

        img = Image.open(io.BytesIO(image_bytes))
        # For ICO, pick the largest size
        if hasattr(img, 'n_frames') and img.format == 'ICO':
            sizes = img.info.get('sizes', set())
            if sizes:
                largest = max(sizes, key=lambda s: s[0] * s[1])
                img.size = largest
                img = img.resize(largest)

        # Convert to RGBA then save as PNG
        if img.mode not in ('RGBA', 'RGB'):
            img = img.convert('RGBA')

        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        img.save(tmp, format='PNG')
        tmp.close()
        return tmp.name
    except Exception as e:
        print(f"  Warning: Could not convert image to PNG: {e}")
        return None


def insert_logo_into_paragraph(paragraph, logo_path, width_inches=2.0):
    """Insert a logo image into a paragraph, replacing existing content.

    Uses python-docx's inline image insertion via run.add_picture().
    """
    clear_paragraph_runs(paragraph)
    run = paragraph.add_run()
    run.add_picture(logo_path, width=Inches(width_inches))


# =====================================================================
# Table of Contents
# =====================================================================

def build_toc_sdt(toc_title_format_para=None):
    """Build a Word TOC as a Structured Document Tag (w:sdt).

    Returns the sdt XML element. The TOC auto-populates when opened in Word.
    """
    sdt = OxmlElement('w:sdt')

    # SDT Properties
    sdtPr = OxmlElement('w:sdtPr')
    docPartObj = OxmlElement('w:docPartObj')
    docPartGallery = OxmlElement('w:docPartGallery')
    docPartGallery.set(qn('w:val'), 'Table of Contents')
    docPartUnique = OxmlElement('w:docPartUnique')
    docPartUnique.set(qn('w:val'), 'true')
    docPartObj.append(docPartGallery)
    docPartObj.append(docPartUnique)
    sdtPr.append(docPartObj)
    sdt.append(sdtPr)

    # SDT Content
    sdtContent = OxmlElement('w:sdtContent')

    # Title paragraph: "Table of Contents"
    title_p = OxmlElement('w:p')
    if toc_title_format_para is not None:
        pPr = toc_title_format_para._element.find(qn('w:pPr'))
        if pPr is not None:
            title_p.insert(0, copy.deepcopy(pPr))

    title_r = OxmlElement('w:r')
    if toc_title_format_para is not None and toc_title_format_para.runs:
        rPr = toc_title_format_para.runs[0]._element.find(qn('w:rPr'))
        if rPr is not None:
            title_r.insert(0, copy.deepcopy(rPr))
    title_t = OxmlElement('w:t')
    title_t.text = 'Table of Contents'
    title_r.append(title_t)
    title_p.append(title_r)
    sdtContent.append(title_p)

    # TOC field code paragraph
    toc_p = OxmlElement('w:p')
    toc_r = OxmlElement('w:r')

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-1" \\h \\z \\u '

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    toc_r.append(fldChar_begin)
    toc_r.append(instrText)
    toc_r.append(fldChar_separate)
    toc_r.append(fldChar_end)
    toc_p.append(toc_r)

    sdtContent.append(toc_p)
    sdt.append(sdtContent)

    return sdt


def set_updatefields_on_open(doc):
    """Set the document to prompt to update fields (including TOC) when opened."""
    settings_element = doc.settings.element
    # Remove existing updateFields if present
    for existing in settings_element.findall(qn('w:updateFields')):
        settings_element.remove(existing)
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings_element.append(update_fields)


# =====================================================================
# Formatting Helpers
# =====================================================================

def clone_run_formatting(source_run, target_run):
    """Copy all formatting from source_run to target_run."""
    rPr = source_run._element.find(qn('w:rPr'))
    if rPr is not None:
        new_rPr = copy.deepcopy(rPr)
        existing = target_run._element.find(qn('w:rPr'))
        if existing is not None:
            target_run._element.remove(existing)
        target_run._element.insert(0, new_rPr)


def set_run_text_with_format(paragraph, text, font_name="Georgia",
                             font_size=None, bold=None, italic=None,
                             color=None):
    """Add a run to a paragraph with specific formatting."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def clear_paragraph_runs(paragraph):
    """Remove all runs from a paragraph (including inline images/drawings)."""
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    # Also remove any remaining r elements
    for r in paragraph._element.findall(qn('w:r')):
        paragraph._element.remove(r)


def replace_paragraph_text(paragraph, new_text, reference_run=None):
    """Replace all text in a paragraph while preserving formatting from the first run."""
    if reference_run is None and paragraph.runs:
        reference_run = paragraph.runs[0]

    saved_rPr = None
    if reference_run is not None:
        rPr = reference_run._element.find(qn('w:rPr'))
        if rPr is not None:
            saved_rPr = copy.deepcopy(rPr)

    clear_paragraph_runs(paragraph)

    run = paragraph.add_run(new_text)
    if saved_rPr is not None:
        existing = run._element.find(qn('w:rPr'))
        if existing is not None:
            run._element.remove(existing)
        run._element.insert(0, saved_rPr)

    return run


def replace_client_name_in_paragraph(paragraph, old_name, new_name):
    """Replace client name across runs in a paragraph, preserving formatting."""
    full_text = paragraph.text
    if old_name not in full_text:
        return

    runs = paragraph.runs
    if not runs:
        return

    ref_rPr = None
    for run in runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            ref_rPr = copy.deepcopy(rPr)
            break

    new_full_text = full_text.replace(old_name, new_name)

    clear_paragraph_runs(paragraph)
    run = paragraph.add_run(new_full_text)
    if ref_rPr is not None:
        existing = run._element.find(qn('w:rPr'))
        if existing is not None:
            run._element.remove(existing)
        run._element.insert(0, ref_rPr)


def create_body_paragraph(doc, text, template_para=None, insert_after=None):
    """Create a new body paragraph with standard formatting."""
    if insert_after is not None:
        new_para = OxmlElement('w:p')

        if template_para is not None:
            pPr = template_para._element.find(qn('w:pPr'))
            if pPr is not None:
                new_pPr = copy.deepcopy(pPr)
                new_para.insert(0, new_pPr)

        r = OxmlElement('w:r')
        if template_para is not None and template_para.runs:
            rPr = template_para.runs[0]._element.find(qn('w:rPr'))
            if rPr is not None:
                r.insert(0, copy.deepcopy(rPr))

        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        new_para.append(r)

        insert_after.addnext(new_para)
        return new_para
    else:
        para = doc.add_paragraph()
        run = para.add_run(text)
        if template_para is not None and template_para.runs:
            clone_run_formatting(template_para.runs[0], run)
            pPr = template_para._element.find(qn('w:pPr'))
            if pPr is not None:
                new_pPr = copy.deepcopy(pPr)
                existing = para._element.find(qn('w:pPr'))
                if existing is not None:
                    para._element.remove(existing)
                para._element.insert(0, new_pPr)
        return para


def create_list_paragraph(text, template_list_para):
    """Create a new list paragraph element cloned from a template list paragraph."""
    new_para = copy.deepcopy(template_list_para._element)
    for r in new_para.findall(qn('w:r')):
        new_para.remove(r)
    for h in new_para.findall(qn('w:hyperlink')):
        new_para.remove(h)

    r = OxmlElement('w:r')
    if template_list_para.runs:
        rPr = template_list_para.runs[0]._element.find(qn('w:rPr'))
        if rPr is not None:
            r.insert(0, copy.deepcopy(rPr))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_para.append(r)
    return new_para


def create_subsection_heading(text, template_heading):
    """Create a new subsection heading element cloned from a template heading."""
    new_para = copy.deepcopy(template_heading._element)
    for r in new_para.findall(qn('w:r')):
        new_para.remove(r)

    r = OxmlElement('w:r')
    if template_heading.runs:
        rPr = template_heading.runs[0]._element.find(qn('w:rPr'))
        if rPr is not None:
            r.insert(0, copy.deepcopy(rPr))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_para.append(r)
    return new_para


def create_body_para_element(text, template_para):
    """Create a body paragraph XML element cloned from a template paragraph."""
    new_para = copy.deepcopy(template_para._element)
    for r in new_para.findall(qn('w:r')):
        new_para.remove(r)
    for h in new_para.findall(qn('w:hyperlink')):
        new_para.remove(h)

    r = OxmlElement('w:r')
    if template_para.runs:
        rPr = template_para.runs[0]._element.find(qn('w:rPr'))
        if rPr is not None:
            r.insert(0, copy.deepcopy(rPr))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_para.append(r)
    return new_para


def remove_element(element):
    """Remove an XML element from its parent."""
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


# =====================================================================
# Table Builders
# =====================================================================

def create_rate_table(rate_data):
    """Create a rate schedule table matching the template format.

    rate_data: {"rows": [["Level", "Rate"], ...]}
    """
    tbl = OxmlElement('w:tbl')

    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    tbl.append(tblPr)

    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(2):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), '3960')
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    header_row = OxmlElement('w:tr')
    for header_text in ["Seniority Level", "Hourly Rate"]:
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '3960')
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '16365C')
        tcPr.append(shd)
        tc.append(tcPr)

        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Georgia')
        rFonts.set(qn('w:hAnsi'), 'Georgia')
        rPr.append(rFonts)
        b = OxmlElement('w:b')
        rPr.append(b)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), 'FFFFFF')
        rPr.append(color)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = header_text
        r.append(t)
        p.append(r)
        tc.append(p)
        header_row.append(tc)
    tbl.append(header_row)

    for level, rate in rate_data["rows"]:
        tr = OxmlElement('w:tr')
        for cell_text in [level, rate]:
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), '3960')
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            tc.append(tcPr)

            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Georgia')
            rFonts.set(qn('w:hAnsi'), 'Georgia')
            rPr.append(rFonts)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = cell_text
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    return tbl


def _make_cell_borders(include_bottom=True):
    """Create tcBorders element matching the template (single, sz=4, black)."""
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        if side == 'bottom' and not include_bottom:
            continue
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    return tcBorders


def _make_cell_margins():
    """Create tcMar element with 100 dxa padding on all sides."""
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'left', 'bottom', 'right']:
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), '100')
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    return tcMar


def _make_timeline_run_props(bold=False, white=False):
    """Create rPr for timeline table cells matching template formatting."""
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Georgia')
    rFonts.set(qn('w:eastAsia'), 'Georgia')
    rFonts.set(qn('w:hAnsi'), 'Georgia')
    rFonts.set(qn('w:cs'), 'Georgia')
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement('w:b'))
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FFFFFF' if white else '000000')
    rPr.append(color)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '21')
    rPr.append(szCs)
    return rPr


def create_timeline_table(table_data):
    """Create a timeline table matching the template format exactly.

    table_data: {"columns": [...], "rows": [{"col_name": "value", ...}, ...]}
    """
    columns = table_data["columns"]
    rows = table_data["rows"]
    num_cols = len(columns)

    tbl = OxmlElement('w:tbl')

    # Table properties - match template exactly
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'a0')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9270')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblJc = OxmlElement('w:jc')
    tblJc.set(qn('w:val'), 'center')
    tblPr.append(tblJc)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    tblLook = OxmlElement('w:tblLook')
    tblLook.set(qn('w:val'), '0400')
    tblLook.set(qn('w:firstRow'), '0')
    tblLook.set(qn('w:lastRow'), '0')
    tblLook.set(qn('w:firstColumn'), '0')
    tblLook.set(qn('w:lastColumn'), '0')
    tblLook.set(qn('w:noHBand'), '0')
    tblLook.set(qn('w:noVBand'), '1')
    tblPr.append(tblLook)
    tbl.append(tblPr)

    # Table grid - distribute widths across 9270 total
    col_widths = []
    base_width = 9270 // num_cols
    remainder = 9270 - (base_width * num_cols)
    for ci in range(num_cols):
        w = base_width + (1 if ci < remainder else 0)
        col_widths.append(str(w))

    tblGrid = OxmlElement('w:tblGrid')
    for cw in col_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), cw)
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # === Header row ===
    header_row = OxmlElement('w:tr')
    trPr = OxmlElement('w:trPr')
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '80')
    trPr.append(trHeight)
    trJc = OxmlElement('w:jc')
    trJc.set(qn('w:val'), 'center')
    trPr.append(trJc)
    header_row.append(trPr)

    for ci, col_name in enumerate(columns):
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), col_widths[ci])
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        tcPr.append(_make_cell_borders())
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0F243E')
        tcPr.append(shd)
        tcPr.append(_make_cell_margins())
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
        tc.append(tcPr)

        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        pPr.append(_make_timeline_run_props(bold=True, white=True))
        p.append(pPr)

        r = OxmlElement('w:r')
        r.append(_make_timeline_run_props(bold=True, white=True))
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = col_name
        r.append(t)
        p.append(r)
        tc.append(p)
        header_row.append(tc)
    tbl.append(header_row)

    # === Data rows ===
    for row_data in rows:
        tr = OxmlElement('w:tr')
        trPr = OxmlElement('w:trPr')
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '889')
        trPr.append(trHeight)
        trJc = OxmlElement('w:jc')
        trJc.set(qn('w:val'), 'center')
        trPr.append(trJc)
        tr.append(trPr)

        # Check if all month columns (1+) have the same value — if so,
        # merge them into a single cell (like the Middlesex template)
        month_values = [row_data.get(col, "") for col in columns[1:]]
        all_same = len(set(month_values)) == 1 and num_cols > 2

        if all_same:
            # --- First cell (row label) ---
            _append_data_cell(tr, row_data.get(columns[0], ""),
                              col_widths[0], is_first_col=True)

            # --- Merged cell spanning all month columns ---
            merged_width = str(sum(int(w) for w in col_widths[1:]))
            merge_span = num_cols - 1
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), merged_width)
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            gridSpan = OxmlElement('w:gridSpan')
            gridSpan.set(qn('w:val'), str(merge_span))
            tcPr.append(gridSpan)
            tcPr.append(_make_cell_borders())
            tcPr.append(_make_cell_margins())
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
            tc.append(tcPr)

            cell_text = month_values[0]
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
            pPr.append(_make_timeline_run_props(bold=False, white=False))
            p.append(pPr)
            r = OxmlElement('w:r')
            r.append(_make_timeline_run_props(bold=False, white=False))
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = cell_text
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        else:
            # --- Individual cells (values differ across months) ---
            for ci, col_name in enumerate(columns):
                cell_text = row_data.get(col_name, "")
                _append_data_cell(tr, cell_text, col_widths[ci],
                                  is_first_col=(ci == 0))
        tbl.append(tr)

    return tbl


def _append_data_cell(tr, cell_text, width, is_first_col=False):
    """Append a single data cell to a timeline table row."""
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), width)
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    tcPr.append(_make_cell_borders())
    tcPr.append(_make_cell_margins())
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)
    tc.append(tcPr)

    lines = cell_text.split('\n') if cell_text else ['']
    for line in lines:
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pBdr = OxmlElement('w:pBdr')
        for side in ['top', 'left', 'bottom', 'right', 'between']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'nil')
            pBdr.append(border)
        pPr.append(pBdr)
        pShd = OxmlElement('w:shd')
        pShd.set(qn('w:val'), 'clear')
        pShd.set(qn('w:color'), 'auto')
        pShd.set(qn('w:fill'), 'FFFFFF')
        pPr.append(pShd)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        pPr.append(_make_timeline_run_props(bold=is_first_col, white=False))
        p.append(pPr)

        r = OxmlElement('w:r')
        r.append(_make_timeline_run_props(bold=is_first_col, white=False))
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = line
        r.append(t)
        p.append(r)
        tc.append(p)
    tr.append(tc)


# =====================================================================
# Compliance Chart Helpers
# =====================================================================

# Default compliance sub-items by customer type
DEFAULT_CONSUMER_COMPLIANCE_ITEMS = [
    "Regulation D (Reserve Requirements)",
    "Regulation DD / Truth in Savings Act",
    "Regulation E (Electronic Fund Transfers)",
    "Regulation CC (Funds Availability)",
    "FDIC Requirements",
]

DEFAULT_COMMERCIAL_COMPLIANCE_ITEMS = [
    "BSA/AML (Bank Secrecy Act)",
    "OFAC Sanctions Compliance",
    "UCC (Uniform Commercial Code)",
    "ACH/Nacha Operating Rules",
    "Commercial Lending Regulations",
]

DEFAULT_BOTH_COMPLIANCE_ITEMS = [
    "BSA/AML & OFAC Sanctions",
    "Regulation E (Electronic Fund Transfers)",
    "Commercial & Consumer Lending Regulations",
    "FDIC & Deposit Compliance",
    "Fair Lending & UDAAP",
]

# Backward-compatible alias
DEFAULT_DEPOSIT_COMPLIANCE_ITEMS = DEFAULT_CONSUMER_COMPLIANCE_ITEMS


def update_compliance_chart(doc, column_header="Deposit Compliance",
                            sub_items=None, customer_type=None):
    """Update Table 0 (compliance programs matrix) third column.

    Replaces the bottom-right header and its sub-items with the provided
    column header and items. When sub_items is None, defaults are chosen
    based on customer_type.
    """
    if sub_items is None:
        if customer_type == "Commercial":
            sub_items = DEFAULT_COMMERCIAL_COMPLIANCE_ITEMS
        elif customer_type == "Both":
            sub_items = DEFAULT_BOTH_COMPLIANCE_ITEMS
        else:
            sub_items = DEFAULT_CONSUMER_COMPLIANCE_ITEMS

    if not doc.tables:
        return

    table = doc.tables[0]  # Table 0: compliance programs matrix

    # Cell [2,2]: header row for the third column (bottom-right of the 2x3 grid)
    if len(table.rows) > 2 and len(table.rows[2].cells) > 2:
        header_cell = table.rows[2].cells[2]
        for para in header_cell.paragraphs:
            if para.text.strip():
                replace_paragraph_text(para, column_header)

    # Cell [3,2]: sub-items for the third column
    if len(table.rows) > 3 and len(table.rows[3].cells) > 2:
        items_cell = table.rows[3].cells[2]

        # Get formatting reference from first paragraph
        ref_para = items_cell.paragraphs[0] if items_cell.paragraphs else None
        ref_rPr = None
        if ref_para and ref_para.runs:
            rPr = ref_para.runs[0]._element.find(qn('w:rPr'))
            if rPr is not None:
                ref_rPr = copy.deepcopy(rPr)

        # Get paragraph properties reference
        ref_pPr = None
        if ref_para is not None:
            pPr = ref_para._element.find(qn('w:pPr'))
            if pPr is not None:
                ref_pPr = copy.deepcopy(pPr)

        # Remove all existing paragraphs from the cell
        for para in items_cell.paragraphs:
            remove_element(para._element)

        # Add new sub-item paragraphs
        tc_elem = items_cell._tc
        for item_text in sub_items:
            p = OxmlElement('w:p')
            if ref_pPr is not None:
                p.insert(0, copy.deepcopy(ref_pPr))
            r = OxmlElement('w:r')
            if ref_rPr is not None:
                r.insert(0, copy.deepcopy(ref_rPr))
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = item_text
            r.append(t)
            p.append(r)
            tc_elem.append(p)


# =====================================================================
# Section Range Helpers
# =====================================================================

def find_section_range(paragraphs, heading_text, end_texts, start_after=0):
    """Find the start and end paragraph indices for a section.

    Returns (heading_idx, end_idx) or (None, None).
    """
    heading_idx = None
    for i, p in enumerate(paragraphs):
        if i < start_after:
            continue
        if p.text.strip() == heading_text:
            heading_idx = i
            break

    if heading_idx is None:
        return None, None

    end_idx = None
    for i, p in enumerate(paragraphs):
        if i > heading_idx and p.text.strip() in end_texts:
            end_idx = i
            break

    return heading_idx, end_idx


def _strip_includepicture_fields(doc):
    """Remove INCLUDEPICTURE field codes from all paragraphs.

    The template contains INCLUDEPICTURE fields pointing to external URLs
    (fsvector.com headshots, etc.). When Word can't reach these URLs, it
    shows broken image placeholder boxes. The actual images are embedded
    as inline drawings in the same runs, so we strip only the field code
    elements (fldChar + instrText) while preserving embedded w:drawing
    elements.
    """
    body = doc.element.body
    # Find all instrText elements with INCLUDEPICTURE anywhere in the doc
    for instr in body.findall('.//' + qn('w:instrText')):
        if instr.text and 'INCLUDEPICTURE' in instr.text:
            # Found an INCLUDEPICTURE field. Remove all fldChar and instrText
            # elements from the parent paragraph while keeping drawings.
            para = instr.getparent()
            # Walk up to find the paragraph element
            while para is not None and para.tag != qn('w:p'):
                para = para.getparent()
            if para is None:
                continue

            # Collect all fldChar and instrText elements in this paragraph
            field_elements = []
            for r in para.findall(qn('w:r')):
                has_drawing = r.find(qn('w:drawing')) is not None
                has_fldChar = r.find(qn('w:fldChar')) is not None
                has_instrText = r.find(qn('w:instrText')) is not None

                if has_drawing:
                    # Keep runs that have embedded drawings
                    # But strip any fldChar/instrText that might also be in this run
                    for fc in r.findall(qn('w:fldChar')):
                        r.remove(fc)
                    for it in r.findall(qn('w:instrText')):
                        r.remove(it)
                elif has_fldChar or has_instrText:
                    # Pure field code run with no drawing - remove entirely
                    field_elements.append(r)

            for elem in field_elements:
                para.remove(elem)


def find_table_in_range(doc, start_elem, end_elem):
    """Find tables positioned between two elements in the document body."""
    body = doc.element.body
    siblings = list(body)
    try:
        start_idx = siblings.index(start_elem)
        end_idx = siblings.index(end_elem)
    except ValueError:
        return []

    tables = []
    for table in doc.tables:
        tbl_elem = table._tbl
        try:
            tbl_idx = siblings.index(tbl_elem)
            if start_idx < tbl_idx < end_idx:
                tables.append(tbl_elem)
        except ValueError:
            continue
    return tables


# =====================================================================
# Main Generator
# =====================================================================

def generate_proposal(
    client_name,
    proposal_type,
    date_str,
    engagement_background_paragraphs,
    engagement_scope_sections,
    budget_timeline_intro,
    budget_timeline_table_data,
    budget_timeline_followup,
    client_website=None,
    include_managed_services=False,
    compliance_column=None,
    compliance_sub_items=None,
    skip_compliance_chart=False,
    customer_type=None,
    output_path=None
):
    """
    Generate a proposal document.

    Args:
        client_name: Name of the client company (e.g., "Acme Corp")
        proposal_type: Title line (e.g., "Regulatory Compliance Advisory Services")
        date_str: Date string (e.g., "February 18, 2026")
        engagement_background_paragraphs: List of strings for Engagement Background
        engagement_scope_sections: List of dicts with:
            - "heading": subsection heading text
            - "intro": intro paragraph text (optional)
            - "bullets": list of bullet point strings (optional)
            - "body_paragraphs": list of additional body paragraph strings (optional)
        budget_timeline_intro: Intro paragraph for Budget and Timeline
        budget_timeline_table_data: Dict with:
            - "columns": list of column header strings
            - "rows": list of dicts with keys matching column headers
        budget_timeline_followup: List of paragraph strings after the table
        client_website: URL of client website for logo fetching (optional)
        include_managed_services: If True, keep Managed Services as primary
            section. If False (default), move to Optional Services.
        compliance_column: Override header for compliance chart column 3
            (default: "Deposit Compliance")
        compliance_sub_items: Override sub-items for compliance chart column 3
        skip_compliance_chart: If True, delete the compliance programs chart
            entirely (for engagements without a compliance program build)
        customer_type: "Commercial", "Consumer", or "Both". Affects default
            compliance sub-items when none are explicitly provided.
        output_path: Where to save. If None, auto-generates path.

    Returns:
        Path to the generated document.
    """
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found at: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    paragraphs = doc.paragraphs

    # ===== 1. UPDATE COVER PAGE =====
    # Paragraph 3: Title
    title_para = paragraphs[3]
    title_text = f"Proposal to Provide {proposal_type}"
    clear_paragraph_runs(title_para)
    run = title_para.add_run(title_text)
    run.font.name = "Georgia"
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x07, 0x37, 0x63)

    # Paragraph 10: Client logo/name area (after "Prepared for:" at P8)
    client_para = paragraphs[10]
    logo_path = None
    if client_website:
        print(f"  Fetching logo from {client_website}...")
        logo_path = fetch_company_logo(client_website)

    if logo_path:
        try:
            insert_logo_into_paragraph(client_para, logo_path, width_inches=2.0)
            print(f"  Logo inserted successfully.")
        except Exception as e:
            print(f"  Warning: Could not insert logo: {e}")
            # Fall back to text
            clear_paragraph_runs(client_para)
            run = client_para.add_run(client_name)
            run.font.name = "Georgia"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        finally:
            try:
                os.unlink(logo_path)
            except OSError:
                pass
    else:
        if client_website:
            print(f"  Warning: Could not fetch logo. Using text instead.")
        clear_paragraph_runs(client_para)
        run = client_para.add_run(client_name)
        run.font.name = "Georgia"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Paragraph 13: Date
    date_para = paragraphs[13]
    replace_paragraph_text(date_para, date_str)

    # ===== 2. UPDATE HEADER =====
    for section in doc.sections:
        header = section.header
        if header.paragraphs:
            header_title = header.paragraphs[0]
            header_text = f"Proposal to Provide {proposal_type}"
            clear_paragraph_runs(header_title)
            run = header_title.add_run(header_text)
            run.font.name = "Georgia"

            if len(header.paragraphs) > 1:
                header_date = header.paragraphs[1]
                try:
                    dt = datetime.strptime(date_str, "%B %d, %Y")
                    month_year = dt.strftime("%B %Y")
                except ValueError:
                    month_year = date_str
                clear_paragraph_runs(header_date)
                run = header_date.add_run(month_year)
                run.font.name = "Georgia"

    # ===== 3. REPLACE TABLE OF CONTENTS =====
    # Remove any existing TOC SDTs from the template
    body = doc.element.body
    existing_toc_sdts = []
    for sdt in body.findall(qn('w:sdt')):
        sdtPr = sdt.find(qn('w:sdtPr'))
        if sdtPr is not None:
            docPartObj = sdtPr.find(qn('w:docPartObj'))
            if docPartObj is not None:
                gallery = docPartObj.find(qn('w:docPartGallery'))
                if gallery is not None and gallery.get(qn('w:val')) == 'Table of Contents':
                    existing_toc_sdts.append(sdt)

    # Find the TOC title paragraph ("Table of Contents")
    toc_title_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'Table of Contents':
            toc_title_idx = i
            break

    if toc_title_idx is not None:
        toc_title_para = doc.paragraphs[toc_title_idx]

        # DON'T scan backward — the empty paragraphs before the TOC title
        # are cover-page padding that keeps content on the title page.
        # Only remove from the TOC title forward.
        toc_start_idx = toc_title_idx

        # Find the end of the TOC area (empty paragraphs + tab after title)
        toc_end_idx = toc_title_idx
        for i in range(toc_title_idx + 1, min(toc_title_idx + 15, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            if p.text.strip() == '' or p.text.strip() == '\t' or p.text == '\t':
                toc_end_idx = i
            else:
                break

        # Build the new TOC SDT
        toc_sdt = build_toc_sdt(toc_title_format_para=toc_title_para)

        # Collect elements to remove BEFORE any insertions (doc.paragraphs
        # indices shift after addprevious, so grab references now).
        anchor = doc.paragraphs[toc_start_idx]._element
        elements_to_remove = []
        for i in range(toc_start_idx, toc_end_idx + 1):
            elements_to_remove.append(doc.paragraphs[i]._element)

        # Insert the TOC SDT right where the old "Table of Contents" text was.
        # The template's empty paragraphs before it already push it to a new page,
        # so we only need a page break AFTER the TOC to separate it from content.
        pb_after = OxmlElement('w:p')
        pb_after_r = OxmlElement('w:r')
        pb_after_br = OxmlElement('w:br')
        pb_after_br.set(qn('w:type'), 'page')
        pb_after_r.append(pb_after_br)
        pb_after.append(pb_after_r)

        # Insert TOC SDT and trailing page break before the old title paragraph
        anchor.addprevious(toc_sdt)
        anchor.addprevious(pb_after)

        # Remove old TOC title + trailing placeholder paragraphs
        for elem in elements_to_remove:
            remove_element(elem)

    # Remove existing template TOC SDTs (the template has one with stale entries)
    for sdt in existing_toc_sdts:
        remove_element(sdt)

    # Set document to update fields on open
    set_updatefields_on_open(doc)

    # ===== 4. UPDATE COMPLIANCE CHART =====
    if skip_compliance_chart:
        # Remove the compliance programs table entirely
        if doc.tables:
            tbl_element = doc.tables[0]._tbl
            # Also remove the paragraph immediately before the table if it's a
            # heading/intro for the chart (e.g., "Compliance Programs")
            prev = tbl_element.getprevious()
            tbl_element.getparent().remove(tbl_element)
    else:
        col_header = compliance_column or "Deposit Compliance"
        update_compliance_chart(doc, column_header=col_header,
                                sub_items=compliance_sub_items,
                                customer_type=customer_type)

    # ===== 5. HANDLE MANAGED SERVICES SECTION =====
    if not include_managed_services:
        # Remove the standalone Managed Services section and its rate table.
        # Move a brief description into Optional Services.
        _move_managed_services_to_optional(doc)
    # If include_managed_services=True, the section stays as-is from the template

    # ===== 6. UPDATE ENGAGEMENT BACKGROUND =====
    # Re-read paragraphs after modifications above
    paragraphs = doc.paragraphs

    bg_heading_idx = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == 'Engagement Background' and p.style.name == 'Heading 1':
            bg_heading_idx = i
            break

    about_fsv_idx = None
    for i, p in enumerate(paragraphs):
        if 'About FS Vector' in p.text and i > (bg_heading_idx or 0):
            about_fsv_idx = i
            break

    if bg_heading_idx is not None and about_fsv_idx is not None:
        ref_body_para = paragraphs[bg_heading_idx + 1]

        elements_to_remove = []
        for i in range(bg_heading_idx + 1, about_fsv_idx):
            elements_to_remove.append(paragraphs[i]._element)

        for elem in elements_to_remove:
            remove_element(elem)

        insert_point = paragraphs[bg_heading_idx]._element
        for para_text in engagement_background_paragraphs:
            new_elem = create_body_para_element(para_text, ref_body_para)
            insert_point.addnext(new_elem)
            insert_point = new_elem

    # ===== 7. UPDATE ENGAGEMENT SCOPE =====
    paragraphs = doc.paragraphs

    scope_heading_idx = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == 'Engagement Scope' and p.style.name == 'Heading 1':
            scope_heading_idx = i
            break

    # Find the next major section after scope
    scope_end_texts = ['Ongoing Advisory Services', 'Headmaster',
                       'Optional Services', 'Budget and Timeline',
                       'Managed Services']
    scope_end_idx = None
    for i, p in enumerate(paragraphs):
        if i > (scope_heading_idx or 0):
            text = p.text.strip()
            if any(text.startswith(t) for t in scope_end_texts):
                scope_end_idx = i
                break

    if scope_heading_idx is not None and scope_end_idx is not None:
        ref_subsection = None
        ref_body = None
        ref_list = None
        for i, p in enumerate(paragraphs):
            if i > scope_heading_idx and i < scope_end_idx:
                if p.style.name == 'Heading 1' and ref_subsection is None:
                    ref_subsection = p
                elif p.style.name == 'Normal' and p.text.strip() and ref_body is None:
                    ref_body = p
                elif p.style.name == 'List Paragraph' and ref_list is None:
                    ref_list = p

        elements_to_remove = []
        for i in range(scope_heading_idx + 1, scope_end_idx):
            elements_to_remove.append(paragraphs[i]._element)

        for elem in elements_to_remove:
            remove_element(elem)

        insert_point = paragraphs[scope_heading_idx]._element
        for section in engagement_scope_sections:
            if ref_subsection:
                heading_elem = create_subsection_heading(
                    section["heading"], ref_subsection)
                insert_point.addnext(heading_elem)
                insert_point = heading_elem

            if section.get("intro") and ref_body:
                intro_elem = create_body_para_element(section["intro"], ref_body)
                insert_point.addnext(intro_elem)
                insert_point = intro_elem

            if section.get("bullets") and ref_list:
                for bullet in section["bullets"]:
                    bullet_elem = create_list_paragraph(bullet, ref_list)
                    insert_point.addnext(bullet_elem)
                    insert_point = bullet_elem

            if section.get("body_paragraphs") and ref_body:
                for bp_text in section["body_paragraphs"]:
                    bp_elem = create_body_para_element(bp_text, ref_body)
                    insert_point.addnext(bp_elem)
                    insert_point = bp_elem

    # ===== 8. UPDATE BUDGET AND TIMELINE =====
    # NO rate table in budget section - only timeline table
    paragraphs = doc.paragraphs

    budget_heading_idx = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == 'Budget and Timeline':
            budget_heading_idx = i
            # Remove pageBreakBefore — no need for a page break after Optional Services
            pPr = p._element.find(qn('w:pPr'))
            if pPr is not None:
                pbb = pPr.find(qn('w:pageBreakBefore'))
                if pbb is not None:
                    pPr.remove(pbb)
            break

    budget_end_idx = None
    for i, p in enumerate(paragraphs):
        if i > (budget_heading_idx or 0):
            text = p.text.strip()
            if text == 'FS Vector Leadership Team':
                budget_end_idx = i
                break

    if budget_heading_idx is not None and budget_end_idx is not None:
        ref_budget_body = None
        for i in range(budget_heading_idx + 1, budget_end_idx):
            p = paragraphs[i]
            if p.style.name == 'Normal' and p.text.strip():
                ref_budget_body = p
                break

        elements_to_remove = []
        for i in range(budget_heading_idx + 1, budget_end_idx):
            elements_to_remove.append(paragraphs[i]._element)

        # Find and remove tables in the budget section
        budget_heading_elem = paragraphs[budget_heading_idx]._element
        budget_end_elem = paragraphs[budget_end_idx]._element
        tables_in_budget = find_table_in_range(doc, budget_heading_elem,
                                               budget_end_elem)

        for elem in elements_to_remove:
            remove_element(elem)
        for elem in tables_in_budget:
            remove_element(elem)

        # Insert new budget content
        insert_point = paragraphs[budget_heading_idx]._element

        # Intro paragraph
        if budget_timeline_intro and ref_budget_body:
            intro_elem = create_body_para_element(
                budget_timeline_intro, ref_budget_body)
            insert_point.addnext(intro_elem)
            insert_point = intro_elem

        # Timeline table only (NO rate table here)
        if budget_timeline_table_data:
            timeline_tbl = create_timeline_table(budget_timeline_table_data)
            insert_point.addnext(timeline_tbl)
            insert_point = timeline_tbl

        # Follow-up paragraphs
        if budget_timeline_followup and ref_budget_body:
            for fp_text in budget_timeline_followup:
                fp_elem = create_body_para_element(fp_text, ref_budget_body)
                insert_point.addnext(fp_elem)
                insert_point = fp_elem

    # ===== 9. REPLACE ALL REMAINING REFERENCES TO OLD CLIENT =====
    old_client = "Parent"
    for para in doc.paragraphs:
        if old_client in para.text:
            replace_client_name_in_paragraph(para, old_client, client_name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_client in para.text:
                        replace_client_name_in_paragraph(
                            para, old_client, client_name)

    # ===== 10. STRIP INCLUDEPICTURE FIELD CODES =====
    # The template has INCLUDEPICTURE fields pointing to external URLs
    # (fsvector.com, etc.) that show broken image boxes in Word when offline.
    # The embedded images are still present, so we strip the field codes
    # and keep only the embedded drawings.
    _strip_includepicture_fields(doc)

    # ===== 11. SAVE =====
    if output_path is None:
        safe_name = re.sub(r'[^\w\s-]', '', client_name).strip().replace(' ', '_')
        date_slug = datetime.now().strftime("%Y%m%d")
        output_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            f"FS_Vector_Proposal_{safe_name}_{date_slug}.docx"
        )

    doc.save(output_path)
    print(f"\nProposal saved to: {output_path}")
    return output_path


# =====================================================================
# Managed Services → Optional Services Logic
# =====================================================================

def _move_managed_services_to_optional(doc):
    """Remove standalone Managed Services section and add brief mention
    under Optional Services.

    When managed services are not a primary offering, they should be listed
    as an optional service rather than a standalone section.
    """
    paragraphs = doc.paragraphs

    # --- Find the Managed Services section ---
    ms_heading_idx = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == 'Managed Services' and p.style.name == 'Heading 1':
            ms_heading_idx = i
            break

    if ms_heading_idx is None:
        return  # Section already removed or not found

    # Find where Managed Services ends (next Heading 1 section)
    ms_end_idx = None
    for i, p in enumerate(paragraphs):
        if i > ms_heading_idx and p.style.name == 'Heading 1':
            # Check this is a major section, not a green subsection
            is_green_subsection = False
            for run in p.runs:
                if run.font.color and run.font.color.rgb == RGBColor(0x3F, 0x81, 0x35):
                    is_green_subsection = True
                    break
            if not is_green_subsection:
                ms_end_idx = i
                break

    if ms_end_idx is None:
        ms_end_idx = len(paragraphs)

    # Remove the Managed Services section paragraphs
    elements_to_remove = []
    for i in range(ms_heading_idx, ms_end_idx):
        elements_to_remove.append(paragraphs[i]._element)

    # Also find and remove the rate table that sits within this section
    if ms_heading_idx < len(paragraphs) and ms_end_idx <= len(paragraphs):
        ms_heading_elem = paragraphs[ms_heading_idx]._element
        ms_end_elem = (paragraphs[ms_end_idx]._element
                       if ms_end_idx < len(paragraphs)
                       else None)
        if ms_end_elem is not None:
            tables_in_ms = find_table_in_range(doc, ms_heading_elem, ms_end_elem)
            for tbl_elem in tables_in_ms:
                remove_element(tbl_elem)

    for elem in elements_to_remove:
        remove_element(elem)

    # --- Update Optional Services section ---
    # Re-read paragraphs after removal
    paragraphs = doc.paragraphs

    opt_heading_idx = None
    for i, p in enumerate(paragraphs):
        if 'Optional Services' in p.text.strip() and p.style.name == 'Heading 1':
            opt_heading_idx = i
            break

    if opt_heading_idx is None:
        return

    opt_para = paragraphs[opt_heading_idx]

    # Rename heading to just "Optional Services"
    if opt_para.text.strip() != 'Optional Services':
        replace_paragraph_text(opt_para, 'Optional Services')

    # Find a reference body paragraph and list paragraph in the Optional section
    ref_body = None
    ref_list = None
    ref_subsection = None
    for i in range(opt_heading_idx + 1, min(opt_heading_idx + 15, len(paragraphs))):
        p = paragraphs[i]
        if p.style.name == 'Heading 1':
            break
        if p.style.name == 'Normal' and p.text.strip() and ref_body is None:
            ref_body = p
        if p.style.name == 'List Paragraph' and ref_list is None:
            ref_list = p

    # Also look for a green subsection heading anywhere in the document for reference
    for p in paragraphs:
        if p.style.name == 'Heading 1':
            for run in p.runs:
                if run.font.color and run.font.color.rgb == RGBColor(0x3F, 0x81, 0x35):
                    ref_subsection = p
                    break
            if ref_subsection:
                break

    # Find the last element of the current Optional Services section content
    # (before the next major heading or end of relevant content)
    insert_after_elem = opt_para._element

    # Find existing content end for Optional Services
    for i in range(opt_heading_idx + 1, len(paragraphs)):
        p = paragraphs[i]
        if p.style.name == 'Heading 1':
            # Check if it's a subsection (green) or new major section
            is_green = False
            for run in p.runs:
                if run.font.color and run.font.color.rgb == RGBColor(0x3F, 0x81, 0x35):
                    is_green = True
                    break
            if not is_green:
                break
        insert_after_elem = p._element

    # Add "Staffing and Recruiting" as a green subsection heading
    # (the existing content is the staffing content, now preceded by
    #  a Managed Services subsection)

    # Insert Managed Services subsection at the beginning of Optional Services
    insert_point = opt_para._element

    if ref_subsection:
        ms_sub_heading = create_subsection_heading(
            "Managed Services", ref_subsection)
        insert_point.addnext(ms_sub_heading)
        insert_point = ms_sub_heading

    if ref_body:
        ms_desc = create_body_para_element(
            "FS Vector offers ongoing managed compliance and risk management "
            "operations as an optional service. This includes collections support, "
            "disputes management, consumer complaints handling, "
            "KYC and sanctions screening, compliance monitoring and testing, "
            "AML/fraud transaction monitoring, and management of risk programs. "
            "FS Vector will provide a dedicated compliance team to manage these "
            "controls and coordinate directly with the partner bank and the client.",
            ref_body)
        insert_point.addnext(ms_desc)
        insert_point = ms_desc

        ms_rate_note = create_body_para_element(
            "Managed services are billed on an hourly basis using FS Vector's "
            "standard rate schedule. Volume-based discounts are available as "
            "the program scales. Please contact FS Vector for detailed pricing.",
            ref_body)
        insert_point.addnext(ms_rate_note)
        insert_point = ms_rate_note

    # Add a "Staffing and Recruiting" subsection heading before the existing
    # staffing content
    if ref_subsection:
        staffing_sub_heading = create_subsection_heading(
            "Staffing and Recruiting", ref_subsection)
        insert_point.addnext(staffing_sub_heading)


# =====================================================================
# Interactive Mode
# =====================================================================

def interactive_mode():
    """Run the generator interactively via command line prompts."""
    print("=" * 60)
    print("  FS Vector Proposal Generator")
    print("=" * 60)
    print()

    # Client details
    client_name = input("Client name (e.g., 'Acme Corp'): ").strip()
    client_website = input("Client website URL (e.g., 'https://www.acme.com', or press Enter to skip): ").strip() or None
    proposal_type = input("Proposal type/title (e.g., 'Regulatory Compliance Advisory Services'): ").strip()
    date_str = input(f"Date (e.g., 'February 18, 2026') [default: today]: ").strip()
    if not date_str:
        date_str = datetime.now().strftime("%B %d, %Y")

    # Managed Services
    ms_input = input("\nInclude Managed Services as a primary service? (y/N): ").strip().lower()
    include_managed = ms_input in ('y', 'yes')

    # Compliance chart
    print("\nCompliance chart - third column defaults to 'Deposit Compliance'.")
    comp_override = input("Override compliance column header? (press Enter for default): ").strip() or None

    # Engagement Background
    print("\n--- Engagement Background ---")
    print("Enter paragraphs for the Engagement Background section.")
    print("Type each paragraph, then press Enter. Type 'DONE' when finished.")
    bg_paragraphs = []
    while True:
        text = input(f"  Paragraph {len(bg_paragraphs) + 1}: ").strip()
        if text.upper() == 'DONE':
            break
        if text:
            bg_paragraphs.append(text)

    # Engagement Scope
    print("\n--- Engagement Scope ---")
    print("Define subsections. For each, provide a heading, intro, and bullet points.")
    print("Type 'DONE' when all subsections are entered.")
    scope_sections = []
    while True:
        heading = input(f"\n  Subsection heading (or 'DONE'): ").strip()
        if heading.upper() == 'DONE':
            break

        intro = input("  Intro paragraph (optional, press Enter to skip): ").strip()
        print("  Bullet points (type each, press Enter. 'DONE' when finished):")
        bullets = []
        while True:
            bullet = input(f"    Bullet {len(bullets) + 1}: ").strip()
            if bullet.upper() == 'DONE':
                break
            if bullet:
                bullets.append(bullet)

        print("  Additional body paragraphs (type each, press Enter. 'DONE' when finished):")
        body_paras = []
        while True:
            bp = input(f"    Paragraph {len(body_paras) + 1}: ").strip()
            if bp.upper() == 'DONE':
                break
            if bp:
                body_paras.append(bp)

        scope_sections.append({
            "heading": heading,
            "intro": intro if intro else None,
            "bullets": bullets if bullets else None,
            "body_paragraphs": body_paras if body_paras else None,
        })

    # Budget and Timeline (no rate table prompts)
    print("\n--- Budget and Timeline ---")
    budget_intro = input("Budget intro paragraph: ").strip()

    print("\nTimeline table columns (comma-separated, e.g., 'Month,1-2,3-4,5+'):")
    cols_str = input("  Columns: ").strip()
    columns = [c.strip() for c in cols_str.split(',')] if cols_str else []

    timeline_rows = []
    if columns:
        print(f"Enter rows (type values for each column, 'DONE' when finished):")
        while True:
            row_label = input(f"\n  {columns[0]} value (or 'DONE'): ").strip()
            if row_label.upper() == 'DONE':
                break
            row_data = {columns[0]: row_label}
            for col in columns[1:]:
                val = input(f"  {col} value: ").strip()
                row_data[col] = val
            timeline_rows.append(row_data)

    timeline_data = {"columns": columns, "rows": timeline_rows} if columns else None

    print("\nFollow-up paragraphs after the budget table ('DONE' when finished):")
    followup = []
    while True:
        fp = input(f"  Paragraph {len(followup) + 1}: ").strip()
        if fp.upper() == 'DONE':
            break
        if fp:
            followup.append(fp)

    # Generate
    output = input("\nOutput file path (press Enter for auto): ").strip() or None
    generate_proposal(
        client_name=client_name,
        proposal_type=proposal_type,
        date_str=date_str,
        engagement_background_paragraphs=bg_paragraphs,
        engagement_scope_sections=scope_sections,
        budget_timeline_intro=budget_intro,
        budget_timeline_table_data=timeline_data,
        budget_timeline_followup=followup,
        client_website=client_website,
        include_managed_services=include_managed,
        compliance_column=comp_override,
        output_path=output,
    )


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--interactive":
        interactive_mode()
    else:
        print("FS Vector Proposal Generator")
        print("Usage:")
        print("  python3 generate_proposal.py --interactive")
        print("  Or import and call generate_proposal() from Python")
        print()
        print(f"Template: {TEMPLATE_PATH}")
        print(f"Template exists: {os.path.exists(TEMPLATE_PATH)}")
